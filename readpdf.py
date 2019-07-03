# -*- coding: utf-8 -*-
"""
Created on Sat Jun  8 15:30:22 2019

@author: Administrator
"""

import os
from io import StringIO
from io import open
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, process_pdf
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

def read_pdf(pdf):
    # resource manager
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    # device
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    process_pdf(rsrcmgr, device, pdf)
    device.close()
    content = retstr.getvalue()
    retstr.close()
    # 获取所有行
    lines = str(content).split("\n")
    return lines
 
 
 #%%
 #coding=utf-8

#%%
article_name = os.listdir('G:\婴儿识别项目\TEMP')
article_name.sort()
i=1
test = Document()
p = test.add_paragraph(u'目录')
for article in article_name:
#    if i<=8 :
#        i=i+1
#        continue
    run = p.add_run(article[:len(article)-4])
    run.font.size = Pt(12)
    run.font.name=u'Arial'
    run.font.color.rgb=RGBColor(0,0,255)
    run.bold=True
    run = p.add_run('\n')
    if __name__ == '__main__':
        with open(article, "rb") as my_pdf:
            lines=read_pdf(my_pdf)
            
            count=0
            for line in lines:
                if count==0:
                    count=1
                    continue
                if line==u'∗':
                    break
                if line=='Contents lists available at ScienceDirect' or line=='' or line=='International Journal of Refrigeration ' or line=='journal homepage: www.elsevier.com/locate/ijrefrig ' or line=='a , ' or line=='b , ' or line=='c , ' or line=='d , ' or line=='e , ':
                    continue
                run = p.add_run(line)
                run.font.size = Pt(12)
                run.font.name=u'Arial'
           #print (lines[0])
            run = p.add_run('\n')
            run = p.add_run(lines[0])
            run.font.size = Pt(12)
            run.font.name=u'Arial'
            run = p.add_run('\n')
            run = p.add_run('\n')
        #text.close()
    test.save(u'C:/Users/Administrator/Desktop/目录.docx')
        #print(read_pdf(my_pdf)[0])
    #%%
article_name = os.listdir('G:\婴儿识别项目\TEMP')
article_name.sort()


article=article_name[2]
#    if i<=8 :
#        i=i+1
#        continue
   
if __name__ == '__main__':
    with open(article, "rb") as my_pdf:
        lines=read_pdf(my_pdf)